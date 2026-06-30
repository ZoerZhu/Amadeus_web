"""Safe workspace file read/list tools for Amadeus agents (Codex-style)."""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any

from .path_guard import (
    DENIED_DIR_NAMES,
    DENIED_FILE_NAMES,
    DENIED_SUFFIXES,
    PathGuard,
    compute_sha256,
)

DEFAULT_MAX_BYTES = 128 * 1024
MAX_BYTES_LIMIT = 1024 * 1024
DEFAULT_MAX_ENTRIES = 200
MAX_ENTRIES_LIMIT = 1000

# Text suffix allowlist (kept here; path_guard only handles deny-lists).
TEXT_SUFFIXES = frozenset({
    ".py", ".pyi", ".ts", ".tsx", ".js", ".jsx", ".mjs", ".cjs",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".md", ".rst", ".txt", ".csv", ".tsv", ".log",
    ".html", ".htm", ".css", ".scss", ".less",
    ".go", ".rs", ".java", ".c", ".h", ".cpp", ".hpp", ".cc",
    ".cs", ".rb", ".php", ".swift", ".kt", ".scala",
    ".sh", ".bash", ".zsh", ".ps1", ".bat", ".cmd",
    ".vue", ".svelte", ".sql", ".gitignore", ".env.example",
    ".lock", ".xml", ".svg",
})

UPLOAD_DOCUMENT_SUFFIXES = frozenset({".pdf", ".doc", ".docx", ".xls", ".xlsx"})
READABLE_DOCUMENT_SUFFIXES = frozenset({".pdf", ".docx", ".xls", ".xlsx"})


def is_text_file(path: Path) -> bool:
    return path.suffix.lower() in TEXT_SUFFIXES


def is_readable_file(path: Path) -> bool:
    return is_text_file(path) or path.suffix.lower() in READABLE_DOCUMENT_SUFFIXES


def extract_readable_text(path: Path, *, max_bytes: int) -> tuple[str, str, bool]:
    if is_text_file(path):
        size = path.stat().st_size
        read_bytes = path.read_bytes()[:max_bytes]
        text, encoding = decode_text(read_bytes)
        return text, encoding, size > len(read_bytes)

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path, max_chars=max_bytes)
    if suffix == ".docx":
        return extract_docx_text(path, max_chars=max_bytes)
    if suffix == ".xlsx":
        return extract_xlsx_text(path, max_chars=max_bytes)
    if suffix == ".xls":
        return extract_xls_text(path, max_chars=max_bytes)
    raise ValueError(f"file_reader cannot read unsupported suffix: {path.suffix or path.name}")


def limit_text(text: str, *, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars].rstrip(), True


def extract_pdf_text(path: Path, *, max_chars: int) -> tuple[str, str, bool]:
    try:
        from pypdf import PdfReader
    except Exception as error:
        raise RuntimeError("读取 PDF 需要安装 pypdf，请运行 pip install pypdf。") from error

    reader = PdfReader(str(path))
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"--- Page {index} ---\n{text.strip()}")
        if sum(len(part) for part in parts) > max_chars:
            break
    text, truncated = limit_text("\n\n".join(parts), max_chars=max_chars)
    return text, "pdf-text", truncated


def extract_docx_text(path: Path, *, max_chars: int) -> tuple[str, str, bool]:
    try:
        from docx import Document
    except Exception as error:
        raise RuntimeError("读取 DOCX 需要安装 python-docx，请运行 pip install python-docx。") from error

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("\t".join(cells))
        if rows:
            parts.append(f"--- Table {table_index} ---\n" + "\n".join(rows))
    text, truncated = limit_text("\n\n".join(parts), max_chars=max_chars)
    return text, "docx-text", truncated


def extract_xlsx_text(path: Path, *, max_chars: int) -> tuple[str, str, bool]:
    try:
        from openpyxl import load_workbook
    except Exception as error:
        raise RuntimeError("读取 XLSX 需要安装 openpyxl，请运行 pip install openpyxl。") from error

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = [f"--- Sheet: {sheet.title} ---"]
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append("\t".join(values).rstrip())
                if sum(len(part) for part in parts) + sum(len(item) for item in rows) > max_chars:
                    break
            parts.append("\n".join(rows))
            if sum(len(part) for part in parts) > max_chars:
                break
    finally:
        workbook.close()
    text, truncated = limit_text("\n\n".join(parts), max_chars=max_chars)
    return text, "xlsx-text", truncated


def extract_xls_text(path: Path, *, max_chars: int) -> tuple[str, str, bool]:
    try:
        import xlrd
    except Exception as error:
        raise RuntimeError("读取 XLS 需要安装 xlrd，请运行 pip install xlrd。") from error

    workbook = xlrd.open_workbook(str(path), on_demand=True)
    parts: list[str] = []
    try:
        for sheet_name in workbook.sheet_names():
            sheet = workbook.sheet_by_name(sheet_name)
            rows: list[str] = [f"--- Sheet: {sheet_name} ---"]
            for row_index in range(sheet.nrows):
                values = [str(sheet.cell_value(row_index, col_index)) for col_index in range(sheet.ncols)]
                if any(value.strip() for value in values):
                    rows.append("\t".join(values).rstrip())
                if sum(len(part) for part in parts) + sum(len(item) for item in rows) > max_chars:
                    break
            parts.append("\n".join(rows))
            if sum(len(part) for part in parts) > max_chars:
                break
    finally:
        workbook.release_resources()
    text, truncated = limit_text("\n\n".join(parts), max_chars=max_chars)
    return text, "xls-text", truncated


def decode_text(data: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-replace"


async def run_file_read(
    *,
    path: str,
    workspace_root: str,
    allowed_external_paths: list[str] | None = None,
    start_line: int | None = None,
    end_line: int | None = None,
    max_bytes: int = DEFAULT_MAX_BYTES,
    include_line_numbers: bool = False,
) -> dict[str, Any]:
    """Read a single file and return content + metadata."""
    guard = PathGuard(workspace_root, allowed_external_paths)
    target = guard.resolve_read(path)
    guard.ensure_not_denied(target)
    if not target.exists() or not target.is_file():
        raise FileNotFoundError(f"file not found: {target}")
    if not is_readable_file(target):
        raise ValueError(f"file type not readable: {target.suffix}")

    max_bytes = max(1, min(int(max_bytes), MAX_BYTES_LIMIT))
    stat = target.stat()
    mtime = _iso_mtime(stat.st_mtime)

    if target.suffix.lower() in READABLE_DOCUMENT_SUFFIXES:
        text, encoding, truncated = extract_readable_text(target, max_bytes=max_bytes)
    else:
        raw = target.read_bytes()
        text, encoding = decode_text(raw[:max_bytes])
        if encoding == "utf-8-sig":
            encoding = "utf-8"
        truncated = len(raw) > max_bytes

    # Normalize CRLF/CR to LF so content and sha256 are platform-stable.
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # sha256 and sizeBytes reflect the decoded, normalized content (before line-range/numbering).
    sha256 = compute_sha256(text)
    size_bytes = len(text.encode("utf-8"))

    total_lines = text.count("\n") + (1 if text and not text.endswith("\n") else 0)
    if start_line is not None or end_line is not None:
        lines = text.splitlines(keepends=True)
        s = max(1, int(start_line or 1))
        e = min(len(lines), int(end_line or len(lines)))
        selected = lines[s - 1:e]
        text = "".join(selected)
    else:
        s, e = 1, total_lines

    if include_line_numbers:
        numbered_lines = []
        for idx, line in enumerate(text.splitlines(keepends=True), start=s):
            prefix = f"{idx}→"
            numbered_lines.append(f"{prefix}{line}" if line.endswith("\n") else f"{prefix}{line}\n")
        text = "".join(numbered_lines)

    return {
        "path": str(target),
        "content": text,
        "sha256": sha256,
        "sizeBytes": size_bytes,
        "mtime": mtime,
        "encoding": encoding,
        "truncated": bool(truncated),
        "startLine": s,
        "endLine": e,
        "totalLineCount": total_lines,
    }


async def run_file_list(
    *,
    path: str,
    workspace_root: str,
    allowed_external_paths: list[str] | None = None,
    recursive: bool = False,
    include_hidden: bool = False,
    max_entries: int = DEFAULT_MAX_ENTRIES,
) -> dict[str, Any]:
    """List a directory and return entries with basic metadata."""
    guard = PathGuard(workspace_root, allowed_external_paths)
    target = guard.resolve_read(path)
    guard.ensure_not_denied(target)
    if not target.exists():
        raise FileNotFoundError(f"path not found: {target}")
    if target.is_file():
        return {
            "path": str(target),
            "type": "file",
            "recursive": recursive,
            "entries": [_entry(target)],
            "skipped": 0,
            "truncated": False,
        }

    max_entries = max(1, min(int(max_entries), MAX_ENTRIES_LIMIT))
    entries: list[dict[str, Any]] = []
    skipped = 0
    truncated = False

    def _visit(dir_path: Path):
        nonlocal skipped, truncated
        try:
            children = sorted(dir_path.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
        except PermissionError:
            skipped += 1
            return
        for child in children:
            if len(entries) >= max_entries:
                truncated = True
                return
            if not include_hidden and child.name.startswith("."):
                skipped += 1
                continue
            if child.is_dir() and child.name in DENIED_DIR_NAMES:
                skipped += 1
                continue
            try:
                guard.ensure_not_denied(child)
            except ValueError:
                skipped += 1
                continue
            entries.append(_entry(child))
            if recursive and child.is_dir():
                _visit(child)

    _visit(target)
    return {
        "path": str(target),
        "type": "directory",
        "recursive": recursive,
        "entries": entries,
        "skipped": skipped,
        "truncated": truncated,
    }


def _entry(path: Path) -> dict[str, Any]:
    try:
        stat = path.stat()
        return {
            "path": str(path),
            "name": path.name,
            "type": "directory" if path.is_dir() else "file",
            "sizeBytes": stat.st_size if path.is_file() else 0,
            "modifiedAt": _iso_mtime(stat.st_mtime),
        }
    except OSError:
        return {"path": str(path), "name": path.name, "type": "unknown", "sizeBytes": 0, "modifiedAt": ""}


def _iso_mtime(ts: float) -> str:
    import datetime
    return datetime.datetime.fromtimestamp(ts, tz=datetime.timezone.utc).isoformat()


async def run_file_reader(args: dict[str, Any]) -> dict[str, Any]:
    """Legacy entry point — returns a migration error. No silent compatibility."""
    return {
        "ok": False,
        "summary": (
            "file_read(action=...) is removed. Migrate to file_read (read a file) or "
            "file_list (list a directory). See the new tool schemas."
        ),
        "data": {"legacyArgs": args},
    }


# --- Legacy path helpers (deprecated; kept for callers that have not migrated
#     to the explicit workspace_root API). New code should use PathGuard directly. ---
_WORKSPACE_PATH_VALUE = os.getenv("AMADEUS_WORKSPACE_PATH", "").strip()
WorkspacePath = (
    Path(_WORKSPACE_PATH_VALUE).expanduser().resolve()
    if _WORKSPACE_PATH_VALUE
    else Path(__file__).resolve().parents[3]
)


def resolve_workspace_path(path_text: str) -> Path:
    """Legacy: resolve a path against the module-level WorkspacePath (workspace-only)."""
    candidate = Path(path_text)
    if candidate.is_absolute():
        resolved = candidate.resolve()
    else:
        resolved = (WorkspacePath / candidate).resolve()
    try:
        resolved.relative_to(WorkspacePath)
    except ValueError as error:
        raise ValueError("file_reader path must stay inside the workspace") from error
    return resolved


def ensure_allowed_path(path: Path) -> None:
    """Legacy: deny-list check using PathGuard against the module-level WorkspacePath."""
    PathGuard(str(WorkspacePath)).ensure_not_denied(path)


def relative_path(path: Path) -> str:
    """Legacy: return path relative to WorkspacePath (forward slashes), or absolute str."""
    try:
        return str(path.resolve().relative_to(WorkspacePath)).replace("\\", "/") or "."
    except ValueError:
        return str(path)
