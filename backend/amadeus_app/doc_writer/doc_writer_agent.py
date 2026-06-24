from __future__ import annotations

import csv
import json
import os
import re
from io import StringIO
from datetime import datetime
from pathlib import Path
from typing import Any, Literal, TypedDict
from zoneinfo import ZoneInfo

import httpx
from langgraph.graph import END, START, StateGraph

from .._common import get_saved_settings_payload
from ..domain import ModelSettings
from ..model_adapter import build_chat_payload
from ..providers import get_provider
from ..runtime_config import effective_model_settings
from ..search.web_search_agent import run_web_search_agent


DocFormat = Literal["md", "txt", "docx", "csv", "xlsx"]
WORKSPACE_PATH_VALUE = os.getenv("AMADEUS_WORKSPACE_PATH", "").strip()
WorkspacePath = Path(WORKSPACE_PATH_VALUE).expanduser().resolve() if WORKSPACE_PATH_VALUE else Path(__file__).resolve().parents[3]
DEFAULT_OUTPUT_DIR = Path("generated_docs")
DEFAULT_SECTIONS = ["背景与目标", "核心内容", "实施步骤", "注意事项", "后续工作"]
RESEARCH_DEFAULT_SECTIONS = ["核心结论", "事实梳理", "规格与能力", "定价与可用性", "事件时间线", "行业影响", "信息来源与不确定性"]
DOC_FORMAT_LABELS: dict[str, str] = {
    "md": "Markdown",
    "txt": "TXT",
    "docx": "DOCX",
    "csv": "CSV",
    "xlsx": "XLSX",
}
DOC_FORMAT_EXTENSIONS: dict[str, str] = {
    "md": ".md",
    "txt": ".txt",
    "docx": ".docx",
    "csv": ".csv",
    "xlsx": ".xlsx",
}
SUPPORTED_SUFFIX_FORMATS: dict[str, DocFormat] = {
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
    ".docx": "docx",
    ".csv": "csv",
    ".xlsx": "xlsx",
}
TEXT_OUTPUT_FORMATS = {"md", "txt", "csv"}
LLM_WRITABLE_FORMATS = {"md", "txt", "docx"}
DOC_WRITER_LLM_MAX_CHARS = int(os.getenv("AMADEUS_DOC_WRITER_LLM_MAX_CHARS", "18000"))
DOC_WRITER_LLM_MAX_TOKENS = int(os.getenv("AMADEUS_DOC_WRITER_LLM_MAX_TOKENS", "6144"))
SEARCH_HINTS = (
    "最新",
    "当前",
    "实时",
    "搜索",
    "查询",
    "资料",
    "来源",
    "引用",
    "官方",
    "文档",
    "research",
    "search",
    "source",
    "reference",
    "latest",
)


class DocWriterPlan(TypedDict, total=False):
    title: str
    format: DocFormat
    audience: str
    tone: str
    sections: list[str]
    use_web_search: bool
    search_query: str
    save: bool
    overwrite: bool
    output_path: str
    file_name: str
    front_matter: bool


class DocWriterState(TypedDict, total=False):
    intent: str
    payload: dict[str, Any]
    raw_arguments: dict[str, Any]
    plan: DocWriterPlan
    research: dict[str, Any]
    outline: list[dict[str, Any]]
    markdown: str
    output: dict[str, Any]
    warnings: list[dict[str, Any]]


def build_doc_writer_graph():
    builder = StateGraph(DocWriterState)
    builder.add_node("planner_agent", planner_agent)
    builder.add_node("research_agent", research_agent)
    builder.add_node("outline_agent", outline_agent)
    builder.add_node("markdown_writer_agent", markdown_writer_agent)
    builder.add_node("file_writer_agent", file_writer_agent)
    builder.add_edge(START, "planner_agent")
    builder.add_edge("planner_agent", "research_agent")
    builder.add_edge("research_agent", "outline_agent")
    builder.add_edge("outline_agent", "markdown_writer_agent")
    builder.add_edge("markdown_writer_agent", "file_writer_agent")
    builder.add_edge("file_writer_agent", END)
    return builder.compile()


DOC_WRITER_GRAPH = None


def get_doc_writer_graph():
    global DOC_WRITER_GRAPH
    if DOC_WRITER_GRAPH is None:
        DOC_WRITER_GRAPH = build_doc_writer_graph()
    return DOC_WRITER_GRAPH


async def run_doc_writer_agent(args: dict[str, Any]) -> dict[str, Any]:
    intent = str(args.get("intent") or args.get("instruction") or args.get("topic") or "").strip()
    final_state = await get_doc_writer_graph().ainvoke(
        {
            "intent": intent,
            "payload": args,
            "raw_arguments": {},
            "warnings": [],
        }
    )
    output = final_state.get("output", {})
    plan = final_state.get("plan", {})
    markdown = final_state.get("markdown", "")
    title = plan.get("title", "Untitled Document")
    doc_format = str(plan.get("format") or "md")
    format_label = DOC_FORMAT_LABELS.get(doc_format, doc_format.upper())
    saved = output.get("saved", False)
    path_text = output.get("path", "")
    if saved:
        summary = f"已生成 {format_label} 文档《{title}》，保存到 {path_text}。"
    else:
        summary = f"已生成 {format_label} 文档《{title}》，未保存到文件。"
    data: dict[str, Any] = {
        "agent": "doc_writer_agent",
        "format": doc_format,
        "title": title,
        "markdown": markdown,
        "saved": saved,
        "outputPath": path_text,
        "absolutePath": output.get("absolutePath", ""),
        "byteCount": output.get("byteCount", len(markdown.encode("utf-8"))),
        "mimeType": output.get("mimeType", ""),
        "usedWebSearch": bool(plan.get("use_web_search")),
        "plan": plan,
        "outline": final_state.get("outline", []),
        "research": final_state.get("research", {}),
        "warnings": final_state.get("warnings", []),
    }
    if "content" in output:
        data["content"] = output.get("content", "")
    return {
        "summary": summary,
        "data": data,
    }


async def planner_agent(state: DocWriterState) -> dict[str, Any]:
    payload = state.get("payload", {})
    intent = state.get("intent", "")
    title = clean_title(str(payload.get("title") or payload.get("topic") or infer_title(intent) or "未命名文档"))
    warnings = list(state.get("warnings", []))
    raw_output_path = str(payload.get("outputPath") or payload.get("output_path") or "")
    raw_file_name = str(payload.get("fileName") or payload.get("file_name") or "")
    explicit_format = payload.get("format") or payload.get("type")
    if explicit_format is not None and str(explicit_format).strip():
        doc_format = normalize_requested_doc_format(explicit_format, warnings)
    else:
        doc_format = infer_doc_format_from_path(raw_file_name) or infer_doc_format_from_path(raw_output_path) or "md"
    explicit_sections = normalize_sections(payload.get("sections")) or infer_sections_from_instruction(intent)
    llm_write = should_use_llm_writer(payload)
    sections = explicit_sections or ([] if llm_write else (RESEARCH_DEFAULT_SECTIONS if is_research_document(intent, payload) else DEFAULT_SECTIONS))
    use_web_search = decide_web_search(payload, intent)
    search_query = str(payload.get("searchQuery") or payload.get("search_query") or payload.get("query") or title).strip()
    save = bool(payload.get("save", True))
    overwrite = bool(payload.get("overwrite", False))
    front_matter = bool(payload.get("frontMatter", payload.get("front_matter", False)))
    plan: DocWriterPlan = {
        "title": title,
        "format": doc_format,
        "audience": str(payload.get("audience") or "通用读者").strip(),
        "tone": str(payload.get("tone") or "清晰、结构化、偏工程说明").strip(),
        "sections": sections,
        "use_web_search": use_web_search,
        "search_query": search_query,
        "save": save,
        "overwrite": overwrite,
        "output_path": raw_output_path,
        "file_name": raw_file_name,
        "front_matter": front_matter,
    }
    return {"plan": plan, "warnings": warnings}


async def research_agent(state: DocWriterState) -> dict[str, Any]:
    plan = state.get("plan", {})
    payload = state.get("payload", {})
    warnings = list(state.get("warnings", []))
    sources = normalize_sources(payload.get("sources"))
    if not plan.get("use_web_search"):
        return {
            "research": {
                "query": "",
                "answer": "",
                "results": sources,
                "warnings": [],
                "skipped": True,
            }
        }

    try:
        search_result = await run_web_search_agent(
            {
                "query": plan.get("search_query", ""),
                "intent": state.get("intent", ""),
                "domains": payload.get("domains", []),
                "excludeDomains": payload.get("excludeDomains", []),
                "maxResults": payload.get("maxResults", 5),
                "fetchContent": payload.get("fetchContent", True),
                "freshness": payload.get("freshness", "any"),
            }
        )
        data = search_result.get("data", {})
        merged_sources = [*sources, *data.get("results", [])]
        return {
            "research": {
                "query": data.get("query", plan.get("search_query", "")),
                "answer": data.get("answer", ""),
                "results": merged_sources,
                "warnings": data.get("warnings", []),
                "skipped": False,
            }
        }
    except Exception as error:
        warnings.append(
            {
                "type": "web_search_failed",
                "message": str(error) or error.__class__.__name__,
            }
        )
        return {
            "research": {
                "query": plan.get("search_query", ""),
                "answer": "",
                "results": sources,
                "warnings": [],
                "skipped": False,
            },
            "warnings": warnings,
        }


async def outline_agent(state: DocWriterState) -> dict[str, Any]:
    plan = state.get("plan", {})
    payload = state.get("payload", {})
    user_points = normalize_string_list(payload.get("keyPoints") or payload.get("key_points"))
    outline: list[dict[str, Any]] = []
    for section in plan.get("sections", DEFAULT_SECTIONS):
        outline.append(
            {
                "heading": section,
                "points": section_points(section, user_points),
            }
        )
    return {"outline": outline}


async def markdown_writer_agent(state: DocWriterState) -> dict[str, Any]:
    plan = state.get("plan", {})
    payload = state.get("payload", {})
    research = state.get("research", {})
    outline = state.get("outline", [])
    warnings = state.get("warnings", [])
    title = plan.get("title", "未命名文档")
    body = str(payload.get("content") or payload.get("body") or payload.get("draft") or "").strip()
    sources = research.get("results", [])
    source_refs = build_source_refs(sources)

    if should_use_llm_writer(payload) and str(plan.get("format") or "md") in LLM_WRITABLE_FORMATS:
        try:
            llm_markdown = await generate_llm_markdown(state, source_refs)
            if llm_markdown.strip():
                return {"markdown": llm_markdown}
        except Exception as error:  # noqa: BLE001
            warnings = [
                *warnings,
                {
                    "type": "llm_writer_failed",
                    "message": f"LLM 写作失败，已回退到结构化规则写作：{error}",
                },
            ]
            if not outline:
                fallback_sections = RESEARCH_DEFAULT_SECTIONS if is_research_document(state.get("intent", ""), payload) else DEFAULT_SECTIONS
                outline = [{"heading": section, "points": section_points(section, [])} for section in fallback_sections]

    lines: list[str] = []

    if plan.get("front_matter"):
        lines.extend(
            [
                "---",
                f"title: {title}",
                f"format: {plan.get('format', 'md')}",
                f"createdAt: {current_time_iso()}",
                f"audience: {plan.get('audience', '通用读者')}",
                "---",
                "",
            ]
        )

    lines.extend(
        [
            f"# {title}",
            "",
            "## 摘要",
            "",
            build_summary(state, source_refs),
            "",
            "## 文档信息",
            "",
            f"- 目标读者：{plan.get('audience', '通用读者')}",
            f"- 写作语气：{plan.get('tone', '清晰、结构化')}",
            f"- 是否使用 Web Search：{'是' if plan.get('use_web_search') else '否'}",
            f"- 生成时间：{current_time_text()}",
            "",
            "## 目录",
            "",
        ]
    )
    for item in outline:
        lines.append(f"- [{item.get('heading', '')}](#{slug_anchor(item.get('heading', ''))})")
    lines.append("")

    if body:
        lines.extend(["## 输入草稿", "", body, ""])

    for item in outline:
        heading = item.get("heading", "")
        lines.extend([f"## {heading}", ""])
        section_lines = research_section_lines(heading, state, source_refs)
        if section_lines:
            lines.extend(section_lines)
        else:
            for point in item.get("points", []):
                lines.append(f"- {point}")
            if heading == "核心内容" and source_refs:
                lines.append(f"- 已参考 {len(source_refs)} 个来源，关键事实应优先以来源列表为准。")
            if heading == "注意事项" and (warnings or research.get("warnings")):
                lines.append("- 检索和生成过程中存在风险提示，见文末“风险提示”。")
        lines.append("")

    if source_refs:
        lines.extend(["## 参考来源", ""])
        for index, source in enumerate(source_refs, start=1):
            title_text = source.get("title") or source.get("url") or f"来源 {index}"
            url = source.get("url", "")
            snippet = source.get("snippet") or source.get("contentSummary") or ""
            if url:
                lines.append(f"{index}. [{title_text}]({url})")
            else:
                lines.append(f"{index}. {title_text}")
            if snippet:
                lines.append(f"   - 摘要：{truncate(snippet, 180)}")
        lines.append("")

    all_warnings = [*warnings, *research.get("warnings", [])]
    if all_warnings:
        lines.extend(["## 风险提示", ""])
        for warning in all_warnings:
            lines.append(f"- {warning.get('message') or warning.get('type')}")
        lines.append("")

    if not source_refs:
        lines.extend(["## 后续可扩展", "", "- 接入 Web Search 或手动来源后，可生成带引用的实质调研正文。", "- 增加 HTML/PDF 等更多导出格式。", "- 增加人工确认后写入任意项目路径。", ""])
    return {"markdown": "\n".join(lines)}


async def file_writer_agent(state: DocWriterState) -> dict[str, Any]:
    plan = state.get("plan", {})
    markdown = state.get("markdown", "")
    doc_format = str(plan.get("format") or "md")
    warnings = list(state.get("warnings", []))
    if not plan.get("save", True):
        output: dict[str, Any] = {
            "saved": False,
            "path": "",
            "byteCount": len(markdown.encode("utf-8")),
            "mimeType": mime_type_for_format(doc_format),
        }
        content = render_text_content(doc_format, state)
        if content is not None:
            output["content"] = content
            output["byteCount"] = len(content.encode("utf-8"))
        elif doc_format not in TEXT_OUTPUT_FORMATS:
            warnings.append(
                {
                    "type": "binary_output_not_returned",
                    "message": f"{DOC_FORMAT_LABELS.get(doc_format, doc_format)} 是二进制文档格式，save=false 时仅返回 Markdown 草稿。",
                }
            )
        return {
            "output": output,
            "warnings": warnings,
        }

    output_path = resolve_output_path(plan, markdown)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists() and not plan.get("overwrite", False):
        output_path = unique_path(output_path)
    artifact = write_output_artifact(doc_format, state, output_path)
    return {
        "output": {
            "saved": True,
            "path": str(output_path.relative_to(WorkspacePath)).replace("\\", "/"),
            "absolutePath": str(output_path),
            "byteCount": artifact.get("byteCount", output_path.stat().st_size),
            "mimeType": artifact.get("mimeType", mime_type_for_format(doc_format)),
            **({"content": artifact["content"]} if "content" in artifact else {}),
        }
    }


def decide_web_search(payload: dict[str, Any], intent: str) -> bool:
    value = payload.get("useWebSearch", payload.get("use_web_search", "auto"))
    if isinstance(value, bool):
        return value
    normalized = str(value).strip().lower()
    if normalized in {"true", "yes", "1", "on"}:
        return True
    if normalized in {"false", "no", "0", "off"}:
        return False
    if payload.get("sources") or payload.get("content") or payload.get("body"):
        return False
    text = " ".join([intent, str(payload.get("title", "")), str(payload.get("topic", ""))]).lower()
    return any(hint in text for hint in SEARCH_HINTS)


def normalize_requested_doc_format(value: Any, warnings: list[dict[str, Any]]) -> DocFormat:
    doc_format = doc_format_from_value(value)
    if doc_format:
        return doc_format
    raw = str(value or "").strip() or "unknown"
    warnings.append(
        {
            "type": "unsupported_format",
            "message": f"不支持的文档格式 {raw}，已改用 Markdown。支持格式：md、txt、docx、csv、xlsx。",
        }
    )
    return "md"


def doc_format_from_value(value: Any) -> DocFormat | None:
    normalized = str(value or "").lower().strip().lstrip(".")
    if normalized == "markdown":
        return "md"
    if normalized in DOC_FORMAT_LABELS:
        return normalized  # type: ignore[return-value]
    if normalized in DOC_FORMAT_EXTENSIONS:
        return normalized  # type: ignore[return-value]
    return None


def infer_doc_format_from_path(path_text: str) -> DocFormat | None:
    suffix = Path(str(path_text or "").strip()).suffix.lower()
    return SUPPORTED_SUFFIX_FORMATS.get(suffix)


def mime_type_for_format(doc_format: str) -> str:
    return {
        "md": "text/markdown; charset=utf-8",
        "txt": "text/plain; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "csv": "text/csv; charset=utf-8",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }.get(doc_format, "application/octet-stream")


def render_text_content(doc_format: str, state: DocWriterState) -> str | None:
    markdown = state.get("markdown", "")
    if doc_format == "md":
        return markdown
    if doc_format == "txt":
        return markdown_to_plain_text(markdown)
    if doc_format == "csv":
        return csv_content_from_state(state)
    return None


def write_output_artifact(doc_format: str, state: DocWriterState, output_path: Path) -> dict[str, Any]:
    text_content = render_text_content(doc_format, state)
    if text_content is not None:
        output_path.write_text(text_content, encoding="utf-8")
        return {
            "byteCount": len(text_content.encode("utf-8")),
            "mimeType": mime_type_for_format(doc_format),
            "content": text_content,
        }
    if doc_format == "docx":
        write_docx_document(state, output_path)
    elif doc_format == "xlsx":
        write_xlsx_workbook(state, output_path)
    else:
        raise ValueError(f"Unsupported doc_writer format: {doc_format}")
    return {
        "byteCount": output_path.stat().st_size,
        "mimeType": mime_type_for_format(doc_format),
    }


def markdown_to_plain_text(markdown: str) -> str:
    output: list[str] = []
    in_front_matter = False
    for index, line in enumerate(markdown.splitlines()):
        stripped = line.strip()
        if index == 0 and stripped == "---":
            in_front_matter = True
            continue
        if in_front_matter:
            if stripped == "---":
                in_front_matter = False
            continue
        if stripped.startswith("```"):
            continue
        text = re.sub(r"^#{1,6}\s+", "", line)
        text = re.sub(r"^\s*[-*+]\s+", "- ", text)
        text = re.sub(r"^\s*(\d+)\.\s+", r"\1. ", text)
        output.append(markdown_inline_to_text(text).rstrip())
    return "\n".join(output).strip() + "\n"


def markdown_inline_to_text(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", lambda match: f"{match.group(1)} ({match.group(2)})", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text


def markdown_body_lines(markdown: str) -> list[str]:
    output: list[str] = []
    in_front_matter = False
    for index, line in enumerate(markdown.splitlines()):
        stripped = line.strip()
        if index == 0 and stripped == "---":
            in_front_matter = True
            continue
        if in_front_matter:
            if stripped == "---":
                in_front_matter = False
            continue
        output.append(line)
    return output


def write_docx_document(state: DocWriterState, output_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("生成 docx 需要安装 python-docx，请运行 pip install python-docx。") from error

    plan = state.get("plan", {})
    markdown = state.get("markdown", "")
    document = Document()
    document.core_properties.title = str(plan.get("title") or "Untitled Document")
    document.core_properties.author = "Amadeus doc_writer_agent"

    for line in markdown_body_lines(markdown):
        stripped = line.strip()
        if not stripped:
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(markdown_inline_to_text(heading_match.group(2)), level=level)
            continue
        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet_match:
            document.add_paragraph(markdown_inline_to_text(bullet_match.group(1)), style="List Bullet")
            continue
        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            document.add_paragraph(markdown_inline_to_text(numbered_match.group(1)), style="List Number")
            continue
        document.add_paragraph(markdown_inline_to_text(stripped))

    document.save(output_path)


def csv_content_from_state(state: DocWriterState) -> str:
    payload = state.get("payload", {})
    tabular = normalize_tabular_rows(payload.get("rows") or payload.get("table") or payload.get("dataset"))
    if tabular:
        headers, rows = tabular
        return csv_text([headers, *rows])

    plan = state.get("plan", {})
    research = state.get("research", {})
    source_refs = build_source_refs(research.get("results", []))
    rows: list[list[Any]] = [["type", "title", "content", "url"]]
    rows.append(["summary", plan.get("title", "未命名文档"), build_summary(state, source_refs), ""])
    for item in state.get("outline", []):
        heading = item.get("heading", "")
        for point in item.get("points", []):
            rows.append(["section", heading, point, ""])
    for source in source_refs:
        rows.append(["source", source.get("title", ""), source.get("snippet", ""), source.get("url", "")])
    for warning in [*state.get("warnings", []), *research.get("warnings", [])]:
        rows.append(["warning", warning.get("type", ""), warning.get("message", ""), ""])
    return csv_text(rows)


def csv_text(rows: list[list[Any]]) -> str:
    buffer = StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    for row in rows:
        writer.writerow([stringify_cell(cell) for cell in row])
    return buffer.getvalue()


def normalize_tabular_rows(value: Any) -> tuple[list[str], list[list[Any]]] | None:
    if not isinstance(value, list) or not value:
        return None
    if all(isinstance(item, dict) for item in value):
        headers: list[str] = []
        for item in value:
            for key in item.keys():
                key_text = str(key)
                if key_text not in headers:
                    headers.append(key_text)
        rows = [[item.get(header, "") for header in headers] for item in value]
        return headers, rows
    if all(isinstance(item, (list, tuple)) for item in value):
        rows = [list(item) for item in value]
        headers = [stringify_cell(cell) for cell in rows[0]]
        return headers, rows[1:]
    return None


def write_xlsx_workbook(state: DocWriterState, output_path: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError as error:
        raise RuntimeError("生成 xlsx 需要安装 openpyxl，请运行 pip install openpyxl。") from error

    plan = state.get("plan", {})
    research = state.get("research", {})
    source_refs = build_source_refs(research.get("results", []))
    workbook = Workbook()
    document_sheet = workbook.active
    document_sheet.title = "Document"
    document_sheet.append(["Field", "Value"])
    document_sheet.append(["Title", plan.get("title", "未命名文档")])
    document_sheet.append(["Format", plan.get("format", "xlsx")])
    document_sheet.append(["Audience", plan.get("audience", "通用读者")])
    document_sheet.append(["Tone", plan.get("tone", "清晰、结构化")])
    document_sheet.append(["Used Web Search", "是" if plan.get("use_web_search") else "否"])
    document_sheet.append(["Generated At", current_time_text()])
    document_sheet.append([])
    document_sheet.append(["Section", "Point"])
    for item in state.get("outline", []):
        heading = item.get("heading", "")
        points = item.get("points", [])
        if not points:
            document_sheet.append([heading, ""])
        for point in points:
            document_sheet.append([heading, point])

    payload = state.get("payload", {})
    tabular = normalize_tabular_rows(payload.get("rows") or payload.get("table") or payload.get("dataset"))
    if tabular:
        headers, rows = tabular
        data_sheet = workbook.create_sheet("Data")
        data_sheet.append(headers)
        for row in rows:
            data_sheet.append([stringify_cell(cell) for cell in row])

    if source_refs:
        source_sheet = workbook.create_sheet("Sources")
        source_sheet.append(["Title", "URL", "Snippet"])
        for source in source_refs:
            source_sheet.append([source.get("title", ""), source.get("url", ""), source.get("snippet", "")])

    warnings = [*state.get("warnings", []), *research.get("warnings", [])]
    if warnings:
        warning_sheet = workbook.create_sheet("Warnings")
        warning_sheet.append(["Type", "Message"])
        for warning in warnings:
            warning_sheet.append([warning.get("type", ""), warning.get("message", "")])

    header_fill = PatternFill("solid", fgColor="E9EEF6")
    for sheet in workbook.worksheets:
        sheet.freeze_panes = "A2"
        for cell in sheet[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        for column in sheet.columns:
            width = min(max(len(stringify_cell(cell.value)) for cell in column) + 2, 60)
            sheet.column_dimensions[column[0].column_letter].width = max(width, 12)

    workbook.save(output_path)


def resolve_output_path(plan: DocWriterPlan, markdown: str) -> Path:
    configured_dir = os.getenv("AMADEUS_DOC_WRITER_OUTPUT_DIR", "").strip()
    default_dir = Path(configured_dir) if configured_dir else DEFAULT_OUTPUT_DIR
    raw_output_path = str(plan.get("output_path") or "").strip()
    title = str(plan.get("title") or "document")
    raw_file_name = str(plan.get("file_name") or "").strip()
    doc_format = str(plan.get("format") or "md")
    extension = DOC_FORMAT_EXTENSIONS.get(doc_format, ".md")
    filename = sanitize_filename(raw_file_name) if raw_file_name else ""
    if not filename:
        filename = f"{slug_filename(title)}-{datetime.now(ZoneInfo('Asia/Shanghai')).strftime('%Y%m%d-%H%M%S')}{extension}"
    elif Path(filename).suffix.lower() not in {extension, ".markdown" if doc_format == "md" else extension}:
        filename = str(Path(filename).with_suffix(extension))
    if infer_doc_format_from_path(raw_output_path):
        candidate = Path(raw_output_path)
    else:
        candidate = Path(raw_output_path) / filename if raw_output_path else default_dir / filename
    if candidate.suffix.lower() not in {extension, ".markdown" if doc_format == "md" else extension}:
        candidate = candidate.with_suffix(extension)
    if candidate.is_absolute():
        resolved = candidate.resolve()
    else:
        resolved = (WorkspacePath / candidate).resolve()
    if not is_relative_to(resolved, WorkspacePath):
        raise ValueError("doc_writer outputPath must stay inside the workspace")
    return resolved


def unique_path(path: Path) -> Path:
    stem = path.stem
    suffix = path.suffix
    for index in range(1, 1000):
        candidate = path.with_name(f"{stem}-{index}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError("Unable to allocate a unique document path")


def infer_title(intent: str) -> str:
    clean = re.sub(r"^(写|编写|生成|创建|起草|write|create)\s*", "", intent.strip(), flags=re.IGNORECASE)
    clean = clean.strip(" ：:，,。.")
    if not clean:
        return "未命名文档"
    return truncate(clean, 80)


def clean_title(title: str) -> str:
    title = re.sub(r"\s+", " ", title).strip()
    return truncate(title, 100) or "未命名文档"


def normalize_sections(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        items = re.split(r"[,，\n;；]+", value)
    elif isinstance(value, list):
        items = [str(item) for item in value]
    else:
        return []
    return [clean_title(item) for item in items if clean_title(item)]


def normalize_string_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        parts = re.split(r"[,，\n;；]+", value)
    elif isinstance(value, list):
        parts = [str(item) for item in value]
    else:
        return []
    return [part.strip() for part in parts if part.strip()]


def normalize_sources(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    sources: list[dict[str, Any]] = []
    for item in value:
        if isinstance(item, dict):
            sources.append(dict(item))
        elif str(item).strip():
            sources.append({"title": str(item).strip(), "url": ""})
    return sources


def should_use_llm_writer(payload: dict[str, Any]) -> bool:
    mode = str(payload.get("writingMode") or payload.get("writing_mode") or "").strip().lower()
    if mode in {"rule", "rules", "structured", "template", "deterministic"}:
        return False
    if mode in {"llm", "advanced", "natural", "free", "pro"}:
        return True
    value = payload.get("llmWrite", payload.get("llm_write", os.getenv("AMADEUS_DOC_WRITER_LLM_ENABLED", "1")))
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() not in {"0", "false", "no", "off", "disabled"}


async def generate_llm_markdown(state: DocWriterState, source_refs: list[dict[str, Any]]) -> str:
    plan = state.get("plan", {})
    payload = state.get("payload", {})
    research = state.get("research", {})
    title = str(plan.get("title") or "未命名文档").strip()
    intent = str(state.get("intent") or payload.get("instruction") or payload.get("topic") or "").strip()
    sections = [str(item).strip() for item in plan.get("sections", []) if str(item).strip()]
    evidence_pack = build_llm_evidence_pack(state, source_refs)
    if not evidence_pack.strip() and not intent:
        return ""

    settings_payload = await get_saved_settings_payload()
    model_settings = ModelSettings.model_validate(settings_payload.get("model") or {})
    provider = get_provider(model_settings.provider_name)
    settings = effective_model_settings(model_settings, provider)
    provider = get_provider(settings.provider_name)
    settings = effective_model_settings(settings, provider)
    if not settings.use_remote:
        raise RuntimeError("remote model is disabled")
    if not provider.compatible:
        raise RuntimeError(f"provider {provider.name} is not compatible")
    if not settings.api_key and provider.name.lower() != "ollama":
        raise RuntimeError("model api key is not configured")

    section_instruction = (
        "用户指定了必须覆盖的主题，请以这些主题为主线组织二级标题，但可以合并、改名或补充必要章节：\n"
        + "\n".join(f"- {section}" for section in sections)
        if sections
        else "用户没有要求固定章节。请根据资料自行设计最适合的 Markdown 结构，不要套用固定模板。"
    )
    body_hint = str(payload.get("content") or payload.get("body") or payload.get("draft") or "").strip()
    prompt = f"""\
请根据给定资料写一份完整的中文 Markdown 文档。

写作目标：
- 标题：{title}
- 原始需求：{intent or "未提供"}
- 目标读者：{plan.get("audience", "通用读者")}
- 语气：{plan.get("tone", "清晰、结构化")}
- {section_instruction}

硬性要求：
1. 输出完整 Markdown 正文，从 `# {title}` 开始，不要包裹在代码块中。
2. 使用自然段写作，不要只罗列来源摘要；每个重要结论后用 [S1]、[S2] 这样的来源编号标注。
3. 严格基于资料写作。资料不足时明确写“不确定/未确认/需要官方核验”，不要编造具体数值。
4. 如果资料之间冲突，单独写“信息冲突与不确定性”。
5. 结尾保留“参考来源”，列出来源编号、标题和 URL。
6. 不要写“后续可扩展”“接入 LLM 写作节点”这类系统实现说明。

已有草稿或补充材料：
{truncate_for_prompt(body_hint, 3000) if body_hint else "无"}

检索摘要：
{truncate_for_prompt(str(research.get("answer") or ""), 4000) or "无"}

证据包：
{evidence_pack}
"""

    messages = [
        {
            "role": "system",
            "content": "你是严谨的研究报告作者，擅长把检索材料转成自然、完整、带来源和不确定性标注的中文 Markdown 报告。",
        },
        {"role": "user", "content": prompt},
    ]
    payload_body = build_chat_payload(settings, provider, "fast", messages, stream=False)
    if "max_completion_tokens" in payload_body:
        payload_body["max_completion_tokens"] = DOC_WRITER_LLM_MAX_TOKENS
    elif "max_tokens" in payload_body:
        payload_body["max_tokens"] = DOC_WRITER_LLM_MAX_TOKENS
    else:
        payload_body["max_tokens"] = DOC_WRITER_LLM_MAX_TOKENS
    if "temperature" in payload_body:
        payload_body["temperature"] = float(os.getenv("AMADEUS_DOC_WRITER_LLM_TEMPERATURE", "0.45"))

    endpoint = f"{settings.base_url.rstrip('/')}/chat/completions"
    headers = {"Content-Type": "application/json"}
    if settings.api_key:
        headers["Authorization"] = f"Bearer {settings.api_key}"
    timeout = float(os.getenv("AMADEUS_DOC_WRITER_LLM_TIMEOUT", "120"))
    async with httpx.AsyncClient(timeout=httpx.Timeout(timeout, connect=20.0)) as client:
        response = await client.post(endpoint, headers=headers, json=payload_body)
    if response.status_code < 200 or response.status_code >= 300:
        raise RuntimeError(f"LLM HTTP {response.status_code}: {response.text[:500]}")
    data = response.json()
    content = extract_chat_content(data)
    markdown = clean_llm_markdown(content, title)
    if len(markdown.strip()) < 200:
        raise RuntimeError("LLM returned too little content")
    return markdown


def build_llm_evidence_pack(state: DocWriterState, source_refs: list[dict[str, Any]]) -> str:
    research = state.get("research", {})
    warnings = [*state.get("warnings", []), *research.get("warnings", [])]
    lines: list[str] = []
    for index, source in enumerate(source_refs[:12], start=1):
        title = str(source.get("title") or f"来源 {index}").strip()
        url = str(source.get("url") or "").strip()
        snippet = str(source.get("snippet") or "").strip()
        lines.append(f"[S{index}] {title}")
        if url:
            lines.append(f"URL: {url}")
        if snippet:
            lines.append(f"摘要: {truncate_for_prompt(snippet, 900)}")
        lines.append("")
    if warnings:
        lines.append("检索风险提示：")
        for warning in warnings[:8]:
            message = str(warning.get("message") or warning.get("type") or "").strip()
            if message:
                lines.append(f"- {message}")
    return truncate_for_prompt("\n".join(lines).strip(), DOC_WRITER_LLM_MAX_CHARS)


def extract_chat_content(data: dict[str, Any]) -> str:
    choices = data.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message") or {}
    return str(message.get("content") or "")


def clean_llm_markdown(content: str, title: str) -> str:
    text = str(content or "").strip()
    fence_match = re.match(r"^```(?:markdown|md)?\s*(.*?)\s*```$", text, flags=re.IGNORECASE | re.DOTALL)
    if fence_match:
        text = fence_match.group(1).strip()
    if not text.startswith("# "):
        text = f"# {title}\n\n{text}"
    return text.strip() + "\n"


def truncate_for_prompt(text: str, limit: int) -> str:
    value = str(text or "").strip()
    if len(value) <= limit:
        return value
    return value[: max(0, limit - 80)].rstrip() + "\n...[内容过长，已截断]"


def is_research_document(intent: str, payload: dict[str, Any]) -> bool:
    text = " ".join(
        [
            intent,
            str(payload.get("title", "")),
            str(payload.get("topic", "")),
            str(payload.get("instruction", "")),
        ]
    ).lower()
    return any(keyword in text for keyword in ("调研", "研究", "报告", "research", "report", "latest", "最新"))


def infer_sections_from_instruction(intent: str) -> list[str]:
    text = str(intent or "")
    match = re.search(r"(?:涵盖|包括|包含|覆盖)[:：]\s*([^。.\n]+)", text)
    if not match:
        return []
    candidates = re.split(r"[、,，;；/]+", match.group(1))
    sections = [clean_title(item) for item in candidates if clean_title(item)]
    return sections[:12]


def section_points(section: str, user_points: list[str]) -> list[str]:
    if user_points and section == "核心内容":
        return user_points
    mapping = {
        "背景与目标": ["说明文档要解决的问题、使用场景和读者预期。", "明确本文档的边界，避免把未实现能力写成已完成能力。"],
        "核心内容": ["按照主题拆分关键概念、设计选择和主要结论。", "对需要事实支撑的内容保留来源或待验证标记。"],
        "实施步骤": ["列出可执行步骤，按依赖顺序组织。", "每一步应包含输入、动作和预期输出。"],
        "注意事项": ["标记风险、假设、限制条件和需要人工确认的环节。", "避免让文档读者误解 Agent 已执行未实际执行的操作。"],
        "后续工作": ["列出下一阶段可扩展能力。", "把未完成能力拆成可验证的小任务。"],
    }
    return mapping.get(section, [f"围绕“{section}”补充结构化内容。", "后续可接入 LLM 写作节点生成更丰富段落。"])


def research_section_lines(section: str, state: DocWriterState, source_refs: list[dict[str, Any]]) -> list[str]:
    research = state.get("research", {})
    warnings = [*state.get("warnings", []), *research.get("warnings", [])]
    answer = str(research.get("answer") or "").strip()
    if not source_refs and not answer:
        return []

    title = str(state.get("plan", {}).get("title") or "调研对象")
    section_key = normalize_section_key(section)
    source_claims = source_claim_lines(source_refs)
    warning_messages = [str(item.get("message") or item.get("type") or "").strip() for item in warnings]
    has_weak_primary = any("官方" in item or "primary" in item.lower() for item in warning_messages)
    content_not_fetched = any("网页正文" in item or "content" in item.lower() for item in warning_messages)

    if section_key in {"核心结论", "摘要结论", "结论"}:
        lines = [
            f"- 本报告检索到 {len(source_refs)} 个候选来源；当前材料主要能支撑“第三方资料正在讨论 {title}”这一事实。",
        ]
        if has_weak_primary:
            lines.append("- 未检索到明显的 Anthropic 官方文档、官方公告、论文或代码仓库来源，因此不能把第三方博客中的型号、价格、发布时间等信息视为已确认事实。")
        if content_not_fetched:
            lines.append("- 部分来源无法读取正文，本报告对这些来源只采用标题与摘要级信息。")
        if source_claims:
            lines.append("- 多个第三方来源声称该主题涉及发布时间、定价、基准、可用性或安全事件，但这些说法需要官方来源进一步核验。")
        return lines

    if section_key in {"模型概述", "事实梳理", "背景与目标", "核心内容"}:
        lines = [
            f"围绕“{title}”，检索结果呈现的是一个由第三方站点传播的模型叙事，而不是官方可核验发布页。",
            "",
        ]
        for index, claim in enumerate(source_claims[:6], start=1):
            lines.append(f"- 来源 {index} 提到：{claim}")
        if answer:
            lines.extend(["", "检索摘要：", ""])
            lines.extend([f"- {line}" for line in split_answer_lines(answer)[:8]])
        return lines

    if section_key in {"技术架构与规格", "规格与能力", "技术规格", "上下文", "架构"}:
        return build_keyword_section(
            source_refs,
            keywords=("architecture", "context", "benchmark", "api", "model", "capacity", "上下文", "架构", "规格", "能力"),
            fallback=[
                "- 现有可读材料没有给出可验证的技术白皮书、模型卡或官方 API 规格。",
                "- 第三方来源提及 architecture、context window、API/model string 等信息，但应在官方文档出现前标为未确认。",
            ],
        )

    if section_key in {"基准测试表现", "性能", "benchmark", "benchmarks"}:
        return build_keyword_section(
            source_refs,
            keywords=("benchmark", "score", "performance", "bench", "swe", "基准", "性能", "测试"),
            fallback=[
                "- 未找到官方 benchmark 表或独立评测数据。",
                "- 第三方来源标题/摘要中出现 benchmarks、performance 等描述，但当前证据不足以比较其真实能力。",
            ],
        )

    if section_key in {"定价策略", "定价与可用性", "价格", "pricing"}:
        return build_keyword_section(
            source_refs,
            keywords=("pricing", "price", "$", "api", "tier", "availability", "access", "定价", "价格", "可用"),
            fallback=[
                "- 未找到官方价格页或 API 计费页。",
                "- 若第三方来源提到价格，需要与 Anthropic Console/API pricing 页面交叉验证后再用于决策。",
            ],
        )

    if section_key in {"fable 5 与 mythos 5 的关系", "mythos", "关系"}:
        return build_keyword_section(
            source_refs,
            keywords=("mythos", "fable", "class", "relationship", "关系"),
            fallback=[
                "- 现有来源多把 Fable 5 与 Mythos 5 放在同一叙事中，但缺少官方体系说明。",
                "- 建议把“Fable 5 属于 Mythos-class”这类说法标注为第三方声称。",
            ],
        )

    if section_key in {"出口管制事件时间线", "事件时间线", "当前状态", "时间线", "安全事件"}:
        lines = build_keyword_section(
            source_refs,
            keywords=("suspended", "ban", "restored", "export", "control", "safety", "government", "June", "暂停", "恢复", "出口", "管制", "安全"),
            fallback=[
                "- 未找到官方安全公告或政府文件。",
                "- 第三方来源提及 suspended/restored/export control 等事件线索，但需要官方公告或监管文件确认。",
            ],
        )
        lines.extend(["", "当前状态判断："])
        lines.append("- 在缺少官方确认的情况下，当前状态应写为“未确认”，不应写为正式发布、暂停或恢复。")
        return lines

    if section_key in {"行业影响分析", "影响分析", "行业影响"}:
        return [
            "- 如果 Fable 5/Mythos 5 相关说法属实，其影响主要会落在高端模型竞争、企业 API 成本、代码/Agent 场景能力和安全合规审查上。",
            "- 但目前来源可信度不足，不能据此做采购、迁移或架构选型决策。",
            "- 更稳妥的动作是持续监测 Anthropic 官方模型列表、价格页、API 文档、系统卡/安全说明，以及主流独立评测机构的数据。",
        ]

    if section_key in {"信息来源与不确定性", "注意事项", "风险提示", "不确定性"}:
        lines = []
        if warning_messages:
            lines.extend([f"- {message}" for message in warning_messages if message])
        else:
            lines.append("- 当前材料未产生自动风险提示，但仍建议优先核验官方来源。")
        lines.append("- 报告中的第三方陈述应作为线索，而非最终事实。")
        lines.append("- 若用于正式决策，需补充 Anthropic 官方页面、模型卡、价格页或可信新闻源。")
        return lines

    if section_key in {"实施步骤", "后续工作"}:
        return [
            "- 到 Anthropic 官方 News、Docs、API pricing、model list 页面核验模型是否存在。",
            "- 检索可信新闻源、监管机构公告和独立评测报告，补足官方/第三方交叉验证。",
            "- 将第三方来源中的发布时间、价格、benchmark、上下文长度等字段整理成证据表，并标记可信等级。",
        ]

    return []


def normalize_section_key(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "").strip().lower())
    text = text.replace("：", ":")
    return text


def source_claim_lines(source_refs: list[dict[str, Any]]) -> list[str]:
    claims: list[str] = []
    for source in source_refs:
        title = str(source.get("title") or "").strip()
        snippet = str(source.get("snippet") or "").strip()
        text = truncate("：".join(item for item in (title, snippet) if item), 260)
        if text:
            claims.append(text)
    return claims


def split_answer_lines(answer: str) -> list[str]:
    lines: list[str] = []
    for raw_line in answer.splitlines():
        line = re.sub(r"^\s*\d+\.\s*", "", raw_line).strip()
        if line:
            lines.append(truncate(line, 220))
    return lines


def build_keyword_section(source_refs: list[dict[str, Any]], *, keywords: tuple[str, ...], fallback: list[str]) -> list[str]:
    hits: list[str] = []
    lowered_keywords = tuple(item.lower() for item in keywords)
    for source in source_refs:
        title = str(source.get("title") or "").strip()
        snippet = str(source.get("snippet") or "").strip()
        combined = f"{title} {snippet}".lower()
        if any(keyword in combined for keyword in lowered_keywords):
            hits.append(f"- {truncate('：'.join(item for item in (title, snippet) if item), 260)}")
    if hits:
        hits.append("")
        hits.append("- 上述内容来自搜索摘要/标题层级，未必代表官方确认；关键参数应继续回查原文和官方文档。")
        return hits[:8]
    return fallback


def build_summary(state: DocWriterState, sources: list[dict[str, Any]]) -> str:
    plan = state.get("plan", {})
    intent = state.get("intent", "")
    title = plan.get("title", "未命名文档")
    research = state.get("research", {})
    warnings = [*state.get("warnings", []), *research.get("warnings", [])]
    if sources:
        risk_note = "；但当前来源存在可信度或正文读取限制，结论需保留不确定性" if warnings else ""
        return f"本文围绕“{title}”整理调研报告，结合 {len(sources)} 个来源提炼主要线索、可核验事实和风险判断{risk_note}。原始需求：{intent or '未提供'}。"
    return f"本文围绕“{title}”整理基础文档草稿。当前草稿主要依据用户输入生成，未使用外部检索。原始需求：{intent or '未提供'}。"


def build_source_refs(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    seen: set[str] = set()
    for source in sources:
        url = str(source.get("url") or "").strip()
        title = str(source.get("title") or url or "未命名来源").strip()
        key = url or title
        if not key or key in seen:
            continue
        seen.add(key)
        output.append(
            {
                "title": title,
                "url": url,
                "snippet": source.get("snippet") or source.get("contentSummary") or "",
            }
        )
    return output[:10]


def sanitize_filename(value: str) -> str:
    name = Path(value).name
    name = re.sub(r"[^\w\u4e00-\u9fff.-]+", "-", name, flags=re.UNICODE).strip("-._")
    return name or "document.md"


def slug_filename(value: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff]+", "-", value.lower(), flags=re.UNICODE).strip("-")
    return truncate(slug, 60).strip("-") or "document"


def slug_anchor(value: str) -> str:
    return re.sub(r"\s+", "-", value.strip().lower())


def truncate(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False)


def current_time_iso() -> str:
    return datetime.now(ZoneInfo("Asia/Shanghai")).isoformat()


def current_time_text() -> str:
    return datetime.now(ZoneInfo("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S %Z")


def is_relative_to(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False
